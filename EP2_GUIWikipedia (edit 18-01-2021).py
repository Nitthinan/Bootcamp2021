#GUI Wikipedia
#GUIWiki.py
import wikipedia

#python to Docx คือการสร้างไฟล์ข้อมูลจากการค้นหาให้อยู่ในไฟล์ word
from docx import Document
def Wikipedia(keyword,lang='th'): #Defalt เป็นภาษาไทย
    # data = wikipedia
    wikipedia.set_lang(lang)

    # summary สำหรับบทความที่สรุป
    # data = wikipedia.summary(keyword)

    # page + content บทความทั้งหน้า
    data2 = wikipedia.page(keyword)
    data2 = data2.content

    doc = Document() # สร้างไฟล์ word ใน python
    doc.add_heading(keyword, 0) # 0 คือขนาดของ heading ในไฟล์ word

    doc.add_paragraph(data2)
    doc.save(keyword+'.docx')
    print('Completed')

# try:
    # Wikipedia('ddddfsfs','en')
# except:
    # print('ERROR')
    
wikipedia.set_lang('th') # เปลี่ยนเป็นภาษาไทย
from tkinter import*
from tkinter import ttk
from tkinter import messagebox

GUI = Tk()
GUI.title('Wikipedia')
GUI.geometry('400x300')
#Config
FONT1=('Angsana New',16)

#คำอธิบาย
L=ttk.Label(GUI,text='ค้นหาบทความ',font=FONT1)
L.pack()

#ช่องค้นหาข้อมูล
v_search = StringVar()
E1 = ttk.Entry(GUI,textvariable = v_search,font=FONT1,width=40)
E1.pack(pady=10)

#ปุ่มค้นหาข้อมุล
def Search():
    keyword=v_search.get() # .get() คือ การดึงข้อมูลเข้ามา
    try: #ใชัดักจับ ERROR
        language = v_radio.get() # th / en / zh
        Wikipedia(keyword,language) #ลองค้นหาดูว่าได้ผลลัพธ์หรือไม่ หากหาเจอให้ผ่านไป
        messagebox.showinfo('Save Complete','Save to DOCX')
    except:
        messagebox.showwarning('Keyword Error','Please Try again') #หาค้นหาแล้วไม่เจอ หรือติดปํญหา ให้แสดงข้อความต่อไปนี้

    #print (wikipedia.search(keyword))
    #result = wikipedia.summary(keyword)
    #print(result)

B1 = ttk.Button(GUI,text='Search',command=Search) #Search ต้องเหมือนกับชื่อ Function
B1.pack(ipadx=20,ipady=10,pady=10) # ขนาดของปุ่มค้นหาข้อมูล

#manu เลือกภาษา
F1 = Frame(GUI)
F1.pack(pady=10)

v_radio = StringVar() # ช่องเก็บข้อมูลภาษา
RB1 = ttk.Radiobutton(F1,text='ภาษาไทย',variable=v_radio,value='th')
RB2 = ttk.Radiobutton(F1,text='English',variable=v_radio,value='en')
RB3 = ttk.Radiobutton(F1,text='China',variable=v_radio,value='zh')
RB1.invoke() # คำสั่งให้ค่าเริ่มต้นเป็นภาษาไทย

RB1.grid(row=0,column=0)
RB2.grid(row=0,column=1)
RB3.grid(row=0,column=2)

GUI.mainloop()
